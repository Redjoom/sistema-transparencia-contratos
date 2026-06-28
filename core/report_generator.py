import os
import datetime
from collections import Counter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_oficio(items, original_filename, output_path, user=None):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    now = datetime.datetime.now()
    folio = f"SIT/VP/{now.year}-{now.month:02d}{now.day:02d}"
    fecha = now.strftime("%d/%m/%Y")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SISTEMA INSTITUCIONAL DE TRANSPARENCIA")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONSTANCIA DE PROTECCIÓN DE DATOS PERSONALES")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph().add_run("").font.size = Pt(6)

    # Metadata table (borderless)
    meta = [
        ("OFICIO NÚMERO:", folio),
        ("FECHA:", fecha),
    ]
    if user:
        meta.append(("ELABORÓ:", f"{user['nombre']} ({user['departamento']})"))
    meta.append(("ASUNTO:", "Protección de datos personales — Versión Pública"))
    meta.append(("DOCUMENTO ORIGINAL:", os.path.basename(original_filename)))

    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(meta):
        cell_l = table.cell(i, 0)
        cell_r = table.cell(i, 1)
        run_l = cell_l.paragraphs[0].add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(9)
        run_r = cell_r.paragraphs[0].add_run(value)
        run_r.font.size = Pt(9)
        # Remove borders
        for cell in [cell_l, cell_r]:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders")
            if tcBorders is not None:
                tcPr.remove(tcBorders)

    doc.add_paragraph().add_run("").font.size = Pt(6)

    # Separate with a line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 72)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph().add_run("").font.size = Pt(6)

    # Section title
    p = doc.add_paragraph()
    run = p.add_run("DETALLE DE ELEMENTOS PROTEGIDOS")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph().add_run("").font.size = Pt(4)

    # Count items per (page, type)
    from collections import Counter
    counter = Counter()
    page_set = set()
    total_items = len(items)
    for page_num, label, text, _ in items:
        counter[(page_num, label)] += 1
        page_set.add(page_num)

    grouped = sorted(counter.items(), key=lambda x: (x[0][0], x[0][1]))

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]
    for idx, text in enumerate(["FOJA", "TIPO", "CANTIDAD"]):
        cell = hdr.cells[idx]
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for (page_num, label), count in grouped:
        row = tbl.add_row()
        vals = [str(page_num), label, str(count)]
        for idx, val in enumerate(vals):
            cell = row.cells[idx]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if idx in (0, 2):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run("").font.size = Pt(6)

    # Summary
    p = doc.add_paragraph()
    run = p.add_run("RESUMEN POR FOJA")
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph().add_run("").font.size = Pt(4)

    for pn in sorted(page_set):
        page_total = sum(c for (p, _), c in counter.items() if p == pn)
        p = doc.add_paragraph()
        run = p.add_run(f"  Foja {pn}: {page_total} elemento(s) protegido(s)")
        run.font.size = Pt(9)

    doc.add_paragraph().add_run("").font.size = Pt(4)
    p = doc.add_paragraph()
    run = p.add_run(f"  TOTAL GENERAL: {total_items} elemento(s) protegido(s)")
    run.bold = True
    run.font.size = Pt(10)

    doc.add_paragraph().add_run("").font.size = Pt(12)

    if user:
        p = doc.add_paragraph()
        run = p.add_run(f"Elaboró: {user['nombre']}")
        run.font.size = Pt(9)
        p = doc.add_paragraph()
        run = p.add_run(f"Departamento: {user['departamento']}")
        run.font.size = Pt(9)

    doc.save(output_path)
    return output_path
